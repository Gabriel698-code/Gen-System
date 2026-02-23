import sqlite3
import json
import re
import os
import sys  
import logging
import time
import shutil
import urllib.request
import threading
import subprocess  
import webview
import uvicorn
import hashlib  

# [ALTERADO] CORREÇÃO DE ENCODING (Para o .EXE não travar no Windows)
if sys.platform.startswith('win'):
    if sys.stdout:
        sys.stdout.reconfigure(encoding='utf-8')
    if sys.stderr:
        sys.stderr.reconfigure(encoding='utf-8')

from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional
from io import BytesIO 

# Framework Web (FastAPI)
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

# Bibliotecas de Documentos (PDF, Word, Excel)
from fpdf import FPDF
from docx import Document 
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.chart import BarChart, PieChart, Reference 
import pandas as pd
import PyPDF2

# SDK de Inteligência Artificial (Google Gemini)
import google.generativeai as genai
from google.api_core.exceptions import ResourceExhausted, TooManyRequests, InternalServerError

# Biblioteca de Busca na Web (Deep Search)
try:
    from duckduckgo_search import DDGS
except ImportError:
    print("[AVISO] Biblioteca 'duckduckgo_search' não instalada. O Deep Search não funcionará.")  # [ALTERADO]
    DDGS = None

# ============================================================================
# 1. CONFIGURAÇÃO GERAL E INICIALIZAÇÃO
# ============================================================================

CONFIG_FILE = "user_config.json"
API_KEY_CLIENTE = None

def carregar_chave():
    
    """
    Tenta carregar a chave de API do arquivo de configuração.
    Procura tanto na pasta local quanto na pasta do executável (se compilado).
    """
    global API_KEY_CLIENTE
    caminhos_busca = [CONFIG_FILE]
    
    # Se estiver rodando como .exe (congelado)
    if getattr(sys, 'frozen', False):
        caminhos_busca.append(os.path.join(os.path.dirname(sys.executable), CONFIG_FILE))
        
    for caminho in caminhos_busca:
        if os.path.exists(caminho):
            try:
                with open(caminho, "r") as f:
                    dados = json.load(f)
                    key = dados.get("api_key")
                    if key:
                        genai.configure(api_key=key)
                        API_KEY_CLIENTE = key
                        print("[INFO] Chave carregada com sucesso de:", caminho)  # [ALTERADO]
                        return True
            except Exception as e:
                print("[ERRO] Erro ao ler chave:", e)  # [ALTERADO]
    return False

# Tenta carregar a chave ao iniciar
carregar_chave()
# [FIX COMPATIBILIDADE]
# Alias para versões antigas que chamam carregar_config()
def carregar_config():
    """
    Carrega configurações do sistema (login, senha, etc).
    NÃO confundir com carregar_chave (API).
    """
    caminhos_busca = [CONFIG_FILE]

    if getattr(sys, 'frozen', False):
        caminhos_busca.append(
            os.path.join(os.path.dirname(sys.executable), CONFIG_FILE)
        )

    for caminho in caminhos_busca:
        if os.path.exists(caminho):
            try:
                with open(caminho, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                return {}

    return {}



# Inicializa o App FastAPI
app = FastAPI()

# Configura CORS (Permite que o HTML local converse com o Python)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- DEFINIÇÃO INTELIGENTE DE PASTAS E DIRETÓRIOS ---
# Esta lógica é crucial para que o programa funcione tanto no VS Code quanto como .EXE
if getattr(sys, 'frozen', False):
    # Se for .exe (PyInstaller)
    CAMINHO_BASE = sys._MEIPASS  # Recursos internos temporários (HTML/CSS)
    DIRETORIO_EXECUCAO = os.path.dirname(sys.executable)  # Onde o .exe está (para salvar arquivos)
else:
    # Se for script Python normal
    CAMINHO_BASE = os.path.dirname(os.path.abspath(__file__))
    DIRETORIO_EXECUCAO = CAMINHO_BASE

# Configuração da Pasta de Documentos (Onde os arquivos gerados serão salvos)
PASTA_DOCS = os.path.join(DIRETORIO_EXECUCAO, "documentos")

# Cria a pasta 'documentos' se ela não existir
if not os.path.exists(PASTA_DOCS):
    try:
        os.makedirs(PASTA_DOCS)
        print("[INFO] Pasta 'documentos' criada automaticamente em:", PASTA_DOCS)  # [ALTERADO]
    except Exception as e:
        print("[ERRO] Erro crítico ao criar pasta documentos:", e)  # [ALTERADO]
        PASTA_DOCS = DIRETORIO_EXECUCAO  # Fallback para a raiz se der erro

print("[INFO] Diretório de Recursos (HTML):", CAMINHO_BASE)  # [ALTERADO]
print("[INFO] Diretório de Salvamento (Docs):", PASTA_DOCS)  # [ALTERADO]

# --- ROTAS PARA SERVIR ARQUIVOS ESTÁTICOS ---

@app.get("/")
async def read_index():
    # Rota raiz abre o index.html
    return FileResponse(os.path.join(CAMINHO_BASE, 'index.html'))

@app.get("/index.html")
async def read_index_explicit():
    # Rota explícita para o index.html
    return FileResponse(os.path.join(CAMINHO_BASE, 'index.html'))

# NOVA ROTA DINÂMICA: Permite abrir nfe_simples.html, contrato.html, etc.
@app.get("/{nome_arquivo}.html")
async def servir_paginas_html(nome_arquivo: str):
    caminho = os.path.join(CAMINHO_BASE, f"{nome_arquivo}.html")
    if os.path.exists(caminho):
        return FileResponse(caminho)
    raise HTTPException(status_code=404, detail="Página não encontrada")

# Monta a pasta estática para servir CSS, JS, Imagens se houver
app.mount("/static", StaticFiles(directory=CAMINHO_BASE), name="static")


# ============================================================================
# 2. SERVIÇOS DE DADOS E INTELIGÊNCIA
# ============================================================================

class ExternalDataService:
    """
    Serviço responsável por buscar cotações de moedas e dados financeiros básicos.
    Possui cache para evitar chamadas excessivas à API.
    """
    _cache = {}
    _ttl = 300  # 5 minutos

    @classmethod
    def get_market_data(cls):
        now = time.time()

        if "market" in cls._cache:
            data, timestamp = cls._cache["market"]
            if now - timestamp < cls._ttl:
                return data, True

        try:
            url = "https://economia.awesomeapi.com.br/last/USD-BRL,BTC-BRL,EUR-BRL"
            with urllib.request.urlopen(url, timeout=5) as response:
                raw = json.loads(response.read().decode())

                info_str = "INDICADORES FINANCEIROS (FONTE: AwesomeAPI):\n"
                info_str += f"- Dólar Comercial: R$ {raw['USDBRL']['bid']} (Atualizado: {raw['USDBRL']['create_date']})\n"
                info_str += f"- Bitcoin: R$ {raw['BTCBRL']['bid']} (Atualizado: {raw['BTCBRL']['create_date']})\n"
                info_str += f"- Euro: R$ {raw['EURBRL']['bid']}\n"
                info_str += "- Taxa Selic (Meta): 10.50% a.a. (Referência)\n"

                cls._cache["market"] = (info_str, now)
                return info_str, True

        except Exception as e:
            print("[ERRO] Erro na API externa:", e)
            return "", False


class DeepSearchService:
    """
    Serviço de Busca Profunda na Web usando DuckDuckGo.
    """
    @staticmethod
    def buscar_viabilidade(termo: str) -> str:
        if not DDGS:
            return ""

        print("[INFO] [DEEP SEARCH] Investigando na web:", termo)

        query = f"{termo} brasil regras dados atualizados 2025"

        try:
            with DDGS() as ddgs:
                results = list(ddgs.text(query, region="br-pt", max_results=3))

                if not results:
                    return ""

                texto = "DADOS RECENTES DA WEB (FONTE: DuckDuckGo):\n"
                for r in results:
                    texto += f"- {r['title']}: {r['body']} (Link: {r['href']})\n"

                return texto

        except Exception as e:
            print("[ERRO] Erro no Deep Search:", e)
            return ""


# ============================================================================
# MOTOR DE DECISÃO CONVERSACIONAL 2.5 (CORRIGIDO)
# ============================================================================

def motor_decisao(texto_usuario: str, contexto: dict | None = None):
    """
    Decide a intenção do usuário:
    - tipo_acao: conversa | gerar_documento | gerar_planilha | analisar_arquivo
    - subtipo: contrato | declaracao | estoque | caixa | grafico | simples
    """

    texto = texto_usuario.lower().strip()
    contexto = contexto or {}

    resultado = {
        "tipo_acao": "conversa",
        "subtipo": None,
        "responder": True,
        "gerar_arquivo": False
    }

    # 1. Identifica comandos explícitos de criação
    verbos_criacao = ["crie", "gere", "faça", "fazer", "monte", "criar", "gerar", "montar", "quero", "preciso"]
    quer_criar = any(v in texto.split() for v in verbos_criacao)

    # 2. Identifica se é apenas uma pergunta exploratória
    eh_pergunta = "?" in texto or texto.startswith(
        ("como", "o que", "qual", "quais", "quando", "por que", "porque", "voce", "você")
    )

    # REGRA DE OURO: Se é uma pergunta e não tem verbo de comando, mantém como conversa!
    if eh_pergunta and not quer_criar:
        return resultado

    # -------------------------------
    # PLANILHAS
    # -------------------------------
    if any(p in texto for p in ["planilha", "excel"]):
        if quer_criar or not eh_pergunta:
            resultado["tipo_acao"] = "gerar_planilha"
            resultado["gerar_arquivo"] = True
            resultado["responder"] = False

            if "estoque" in texto:
                resultado["subtipo"] = "estoque"
            elif any(p in texto for p in ["caixa", "fluxo"]):
                resultado["subtipo"] = "caixa"
            elif any(p in texto for p in ["preço", "precificação", "precificacao"]):
                resultado["subtipo"] = "precificacao"
            elif any(p in texto for p in ["gráfico", "grafico"]):
                resultado["subtipo"] = "grafico"
            else:
                resultado["subtipo"] = "simples"

    # -------------------------------
    # DOCUMENTOS (WORD / PDF)
    # -------------------------------
    elif any(p in texto for p in ["contrato", "declaração", "declaraçao", "recibo", "orçamento", "orcamento", "ordem de serviço", "os"]):
        if quer_criar or not eh_pergunta:
            resultado["tipo_acao"] = "gerar_documento"
            resultado["gerar_arquivo"] = True
            resultado["responder"] = False

            if "contrato" in texto:
                resultado["subtipo"] = "contrato"
            elif "declara" in texto:
                resultado["subtipo"] = "declaracao"
            elif "recibo" in texto:
                resultado["subtipo"] = "recibo"
            elif any(p in texto for p in ["orçamento", "orcamento"]):
                resultado["subtipo"] = "orcamento"
            elif any(p in texto for p in ["ordem", "os"]):
                resultado["subtipo"] = "os"

    return resultado

# ============================================================================
# 3. PROCESSAMENTO DE ARQUIVOS E IA (ROUTER)
# ============================================================================

async def ler_arquivo_para_texto(arquivo: UploadFile) -> dict:
    """
    Lê arquivos enviados pelo usuário e converte para texto ou binário.
    Suporta: Imagens, Excel, Word, PDF.
    """
    nome = arquivo.filename.lower()
    conteudo_bytes = await arquivo.read()

    # -------------------------------
    # IMAGENS (MULTIMODAL)
    # -------------------------------
    if nome.endswith((".png", ".jpg", ".jpeg", ".webp")):
        return {
            "tipo": "imagem",
            "conteudo": conteudo_bytes,
            "mime": arquivo.content_type
        }

    # -------------------------------
    # EXCEL
    # -------------------------------
    elif nome.endswith((".xlsx", ".xls")):
        try:
            df = pd.read_excel(BytesIO(conteudo_bytes))
            texto_dados = df.head(50).to_csv(index=False)
            return {
                "tipo": "texto",
                "conteudo": f"DADOS DA PLANILHA (AMOSTRA):\n{texto_dados}",
                "mime": "text/plain"
            }
        except Exception as e:
            return {"tipo": "erro", "conteudo": str(e)}

    # -------------------------------
    # WORD
    # -------------------------------
    elif nome.endswith(".docx"):
        try:
            doc = Document(BytesIO(conteudo_bytes))
            texto = "\n".join(p.text for p in doc.paragraphs)
            return {
                "tipo": "texto",
                "conteudo": f"CONTEÚDO DO DOCUMENTO:\n{texto}",
                "mime": "text/plain"
            }
        except Exception as e:
            return {"tipo": "erro", "conteudo": str(e)}

    # -------------------------------
    # PDF
    # -------------------------------
    elif nome.endswith(".pdf"):
        try:
            reader = PyPDF2.PdfReader(BytesIO(conteudo_bytes))
            texto = ""
            for page in reader.pages:
                texto += (page.extract_text() or "") + "\n"
            return {
                "tipo": "texto",
                "conteudo": f"CONTEÚDO DO PDF:\n{texto}",
                "mime": "text/plain"
            }
        except Exception as e:
            return {"tipo": "erro", "conteudo": str(e)}

    return {"tipo": "erro", "conteudo": "Formato não suportado."}


# ---------------------------------------------------------------------
# MODELOS DE IA (ROUTER COM FALLBACK)
# ---------------------------------------------------------------------

FAST_MODELS = [
    "models/gemini-3-flash-preview",
    "models/gemini-2.5-flash",
    "models/gemini-2.0-flash",
    "models/gemini-flash-latest",
]

COOLDOWN = 120
estado_modelos = {m: {"bloqueado_ate": 0} for m in FAST_MODELS}


def gerar_com_router(
    prompt: str,
    imagem_bytes: Optional[bytes] = None,
    mime_type: str = "image/jpeg"
) -> str:
    """
    Gerenciador inteligente de chamadas à IA.
    Faz fallback automático e suporta multimodal.
    """

    if not API_KEY_CLIENTE:
        return "ERRO: Nenhuma chave de API configurada."

    agora = time.time()
    print("[INFO] [ROUTER] Iniciando processamento IA...")

    conteudo = [prompt]

    if imagem_bytes:
        conteudo.append({
            "mime_type": mime_type,
            "data": imagem_bytes
        })

    for modelo_nome in FAST_MODELS:
        estado = estado_modelos[modelo_nome]

        if agora < estado["bloqueado_ate"]:
            continue

        try:
            model = genai.GenerativeModel(model_name=modelo_nome)
            resp = model.generate_content(
                conteudo,
                request_options={"timeout": 60}
            )
            return resp.text

        except (ResourceExhausted, TooManyRequests):
            estado["bloqueado_ate"] = agora + COOLDOWN
            continue

        except Exception as e:
            print(f"[ERRO] {modelo_nome}:", e)
            continue

    return "⚠️ O sistema está sobrecarregado. Tente novamente em instantes."


# ============================================================================
# EXECUTOR DE DECISÃO (PONTE MOTOR → GERADORES → IA)
# ============================================================================

def executar_decisao_ia(
    texto_usuario: str,
    session_id: str,
    contexto_extra: dict | None = None
) -> dict:
    """
    Integra:
    - Motor de decisão
    - Geradores de arquivos
    - IA conversacional
    """

    contexto_extra = contexto_extra or {}
    decisao = motor_decisao(texto_usuario, contexto_extra)

    tipo = decisao.get("tipo_acao")
    subtipo = decisao.get("subtipo")
    dados = contexto_extra.get("dados_extraidos", {})

    # -------------------------------
    # PLANILHAS
    # -------------------------------
    if tipo == "gerar_planilha":
        if subtipo == "estoque":
            nome = criar_excel_estoque(session_id)
        elif subtipo == "caixa":
            nome = criar_excel_caixa(session_id)
        elif subtipo == "precificacao":
            nome = criar_excel_precificacao(dados, session_id)
        elif subtipo == "grafico":
            nome = criar_excel_com_grafico(dados, session_id)
        else:
            nome = criar_excel_simples(dados, "planilha", session_id)

        # ADICIONADO: Link HTML forçando o download pelo navegador externo
        link_html = f"<a href='http://127.0.0.1:8000/baixar_doc/{nome}' target='_blank' style='color: #8257e5; font-weight: bold;'>📥 CLIQUE AQUI PARA BAIXAR O ARQUIVO</a>"

        return {
            "resposta_usuario": f"📊 **Planilha criada com sucesso!**\n\n{link_html}",
            "arquivo": nome
        }

    # -------------------------------
    # DOCUMENTOS
    # -------------------------------
    if tipo == "gerar_documento":
        if subtipo == "contrato":
            nome = criar_word("contrato", dados, session_id)
        elif subtipo == "declaracao":
            nome = criar_word_declaracao(dados, session_id)
        elif subtipo == "os":
            nome = criar_word_os(dados, session_id)
        elif subtipo == "recibo":
            nome = criar_pdf("recibo", dados, session_id)
        elif subtipo == "orcamento":
            nome = criar_pdf("orcamento", dados, session_id)
        else:
            return {"resposta_usuario": "Tipo de documento não reconhecido."}

        # ADICIONADO: Link HTML forçando o download pelo navegador externo
        link_html = f"<a href='http://127.0.0.1:8000/baixar_doc/{nome}' target='_blank' style='color: #8257e5; font-weight: bold;'>📄 CLIQUE AQUI PARA BAIXAR O ARQUIVO</a>"

        return {
            "resposta_usuario": f"📄 **Documento criado com sucesso!**\n\n{link_html}",
            "arquivo": nome
        }

    # -------------------------------
    # CONVERSA NORMAL
    # -------------------------------
    resposta = gerar_com_router(texto_usuario)
    return {"resposta_usuario": resposta}
# ======================================================================
# CONSTANTES DE PROMPT (BASE / MODOS ESPECIALIZADOS)
# ======================================================================

BASE_JSON_INSTRUCT = """
Responda APENAS em JSON válido.

REGRAS OBRIGATÓRIAS:
1. O campo 'resposta_usuario' DEVE existir e pode usar Markdown.
2. NÃO escreva texto fora do JSON.
3. NÃO invente links ou fontes.
4. Se identificar dados estruturados, preencha 'dados_extraidos'.
5. Se identificar solicitação de documento ou planilha, preencha 'documento_solicitado'.
6. Se houver lista de valores, use 'dados_grafico'.

FORMATO ESPERADO:
{
  "resposta_usuario": "Texto em Markdown",
  "dados_extraidos": {},
  "documento_solicitado": null,
  "dados_grafico": []
}
"""

PROMPTS_MODOS = {

    # ------------------------------------------------------------------
    # MODO GERAL
    # ------------------------------------------------------------------
    "geral": f"""
CONTEXTO:
Você é o Gen, um assistente empresarial inteligente e consultivo.

OBJETIVO:
Ajudar o usuário a entender, planejar e executar decisões de negócio.

DIRETRIZES:
- Linguagem clara e acessível
- Estruture a resposta em tópicos quando possível
- Não assuma informações não fornecidas

{BASE_JSON_INSTRUCT}
""",

    # ------------------------------------------------------------------
    # MODO JURÍDICO
    # ------------------------------------------------------------------
    "juridico": f"""
CONTEXTO:
Você é o Gen Jurídico, especialista em legislação brasileira aplicada a negócios.

FONTES (PRIORIDADE):
1. Deep Search (leis, normas e decisões MAIS RECENTES).
2. Banco de Dados local (leads.db) para conceitos consolidados.
→ Em caso de divergência, PRIORIZE A WEB.

DIRETRIZES:
- Explique implicações legais de forma prática
- Cite riscos, obrigações e cuidados
- NÃO forneça aconselhamento ilegal ou definitivo
- Utilize linguagem clara, não excessivamente técnica

{BASE_JSON_INSTRUCT}
""",

    # ------------------------------------------------------------------
    # MODO FINANCEIRO
    # ------------------------------------------------------------------
    "financeiro": f"""
CONTEXTO:
Você é o Gen Financeiro, analista de finanças empresariais.

FONTES (PRIORIDADE):
1. Deep Search (leis fiscais, regras tributárias, índices atualizados).
2. Banco de Dados local (leads.db) para CNAEs, Simples Nacional e faixas.
→ Se houver conflito, PRIORIZE DADOS DA WEB.

DIRETRIZES:
- Faça cálculos quando possível
- Explique impostos, custos, margens e riscos
- Use exemplos práticos
- Seja conservador nas estimativas

{BASE_JSON_INSTRUCT}
""",

    # ------------------------------------------------------------------
    # MODO MARKETING
    # ------------------------------------------------------------------
    "marketing": f"""
CONTEXTO:
Você é o Gen Marketing, estrategista de crescimento e posicionamento.

OBJETIVO:
Criar estratégias de marketing viáveis para o contexto do cliente.

DIRETRIZES:
- Defina público-alvo
- Sugira canais (online/offline)
- Apresente métricas (CAC, ROI, conversão)
- Traga ideias práticas e executáveis
- Evite promessas irreais

{BASE_JSON_INSTRUCT}
""",

    # ------------------------------------------------------------------
    # MODO VIABILIDADE
    # ------------------------------------------------------------------
    "viabilidade": f"""
CONTEXTO:
Você é o Gen Analista de Viabilidade de Negócios.

PROCESSO OBRIGATÓRIO (5 ETAPAS):
1. Análise do mercado local e concorrência (Deep Search)
2. Avaliação do modelo de negócio
3. Custos, receitas e riscos
4. Análise comparativa (negócios semelhantes / região)
5. Veredito final (Viável / Viável com ajustes / Não viável)

FONTES:
- PRIORIZE Deep Search para dados de mercado
- Use banco local apenas como apoio técnico

DIRETRIZES:
- Seja honesto e técnico
- Aponte riscos reais
- Não incentive negócios inviáveis

{BASE_JSON_INSTRUCT}
"""
}

#--------------------------------------------------------------------------
# -----------------------------------------------------------------------
# ============================================================================
# 4. MODELOS DE DADOS (PYDANTIC) E BANCO DE DADOS
# ============================================================================

DB_FILE = "leads.db"
logging.basicConfig(level=logging.ERROR)

# ---------------------------------------------------------------------
# MODELOS PYDANTIC
# ---------------------------------------------------------------------

class ConfigData(BaseModel):
    api_key: str


class Pedido(BaseModel):
    session_id: str
    texto: str
    modo: Optional[str] = "geral"


class DadosFormulario(BaseModel):
    session_id: str
    tipo: str
    formato: Optional[str] = "pdf"
    dados: Dict[str, Any]


class NotaFiscal(BaseModel):
    numero: int
    serie: int
    chave: str
    emit_cnpj: str
    dest_doc: str
    dest_nome: str
    valor: float
    tipo: str
    xml: str


class OrcamentoData(BaseModel):
    session_id: str
    cliente_nome: str
    cliente_doc: str
    valor_total: float
    validade: str
    itens_json: str


# ---------------------------------------------------------------------
# BANCO DE DADOS
# ---------------------------------------------------------------------

def get_db():
    conn = sqlite3.connect(
        DB_FILE,
        timeout=10,
        check_same_thread=False
    )
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    conn = get_db()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS leads (
            id INTEGER PRIMARY KEY,
            session_id TEXT,
            data_registro TEXT,
            ramo TEXT,
            estagio TEXT,
            capital TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS clientes (
            id INTEGER PRIMARY KEY,
            session_id TEXT,
            data_cadastro TEXT,
            nome TEXT,
            tipo_negocio TEXT,
            documento_tipo TEXT,
            documento_numero TEXT,
            email TEXT,
            investimento TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS mensagens (
            id INTEGER PRIMARY KEY,
            session_id TEXT,
            role TEXT,
            content TEXT,
            timestamp TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS documentos (
            id INTEGER PRIMARY KEY,
            session_id TEXT,
            nome_arquivo TEXT,
            tipo TEXT,
            criado_em TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS sessoes (
            session_id TEXT PRIMARY KEY,
            titulo TEXT,
            criada_em TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS notas_fiscais_clientes (
            id INTEGER PRIMARY KEY,
            numero_nota INTEGER,
            serie INTEGER,
            chave_acesso TEXT,
            emitente_cnpj TEXT,
            destinatario_doc TEXT,
            destinatario_nome TEXT,
            valor_total REAL,
            data_emissao TEXT,
            tipo_nota TEXT,
            xml_completo TEXT,
            criado_em TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS orcamentos_clientes (
            id INTEGER PRIMARY KEY,
            session_id TEXT,
            cliente_nome TEXT,
            cliente_doc TEXT,
            valor_total REAL,
            data_emissao TEXT,
            validade TEXT,
            itens_json TEXT,
            criado_em TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS tabela_cnaes (
            codigo TEXT,
            descricao TEXT,
            anexo_simples TEXT,
            aliquota_inicial TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS tabela_simples (
            anexo TEXT,
            faixa INTEGER,
            limite_faturamento REAL,
            aliquota REAL,
            deducao REAL
        )
    """)

    conn.commit()
    conn.close()
    popular_tabelas_iniciais()


# ---------------------------------------------------------------------
# POPULAÇÃO INICIAL (FIX: popular_tabelas_iniciais)
# ---------------------------------------------------------------------

def popular_tabelas_iniciais():
    conn = get_db()
    cur = conn.cursor()

    if not cur.execute("SELECT COUNT(*) FROM tabela_cnaes").fetchone()[0]:
        cnaes = [
            ("4781-4/00", "Comércio de vestuário", "Anexo I", "4.0%"),
            ("6201-5/00", "Desenvolvimento de software", "Anexo III", "6.0%"),
            ("7319-0/02", "Marketing", "Anexo III", "6.0%"),
        ]
        cur.executemany("INSERT INTO tabela_cnaes VALUES (?,?,?,?)", cnaes)

    if not cur.execute("SELECT COUNT(*) FROM tabela_simples").fetchone()[0]:
        simples = [
            ("Anexo I", 1, 180000, 4.0, 0),
            ("Anexo III", 1, 180000, 6.0, 0),
        ]
        cur.executemany("INSERT INTO tabela_simples VALUES (?,?,?,?,?)", simples)

    conn.commit()
    conn.close()


# ---------------------------------------------------------------------
# UTILIDADES DE CONVERSA / HISTÓRICO (FIXES)
# ---------------------------------------------------------------------

def salvar_mensagem(session_id, role, content):
    conn = get_db()
    cur = conn.cursor()

    cur.execute(
        "INSERT INTO mensagens (session_id, role, content, timestamp) VALUES (?, ?, ?, ?)",
        (session_id, role, content, datetime.now().isoformat())
    )

    if role == "user":
        existe = cur.execute(
            "SELECT 1 FROM sessoes WHERE session_id = ?",
            (session_id,)
        ).fetchone()

        if not existe:
            cur.execute(
                "INSERT INTO sessoes (session_id, titulo, criada_em) VALUES (?, ?, ?)",
                (session_id, content[:30], datetime.now().isoformat())
            )

    conn.commit()
    conn.close()


def get_historico_db(session_id: str) -> str:
    conn = get_db()
    rows = conn.execute(
        "SELECT role, content FROM mensagens WHERE session_id = ? ORDER BY id ASC",
        (session_id,)
    ).fetchall()
    conn.close()

    return "\n".join([f"{r['role']}: {r['content']}" for r in rows])


def buscar_dados_tecnicos(texto_usuario: str) -> str:
    texto = texto_usuario.lower()
    info = ""

    conn = get_db()
    cur = conn.cursor()

    if "cnae" in texto:
        res = cur.execute(
            "SELECT codigo, descricao FROM tabela_cnaes LIMIT 3"
        ).fetchall()
        for r in res:
            info += f"\n- CNAE {r['codigo']}: {r['descricao']}"

    conn.close()
    return info


def salvar_documento_db(session_id, nome_arquivo, tipo):
    conn = get_db()
    conn.execute(
        "INSERT INTO documentos (session_id, nome_arquivo, tipo, criado_em) VALUES (?, ?, ?, ?)",
        (session_id, nome_arquivo, tipo, datetime.now().isoformat())
    )
    conn.commit()
    conn.close()


def formatar_valor(valor_raw):
    if valor_raw in (None, "", {}):
        return "R$ 0,00"

    try:
        if isinstance(valor_raw, (int, float)):
            val = float(valor_raw)
        else:
            limpo = re.sub(r"[^\d.,-]", "", str(valor_raw))
            val = float(limpo.replace(".", "").replace(",", "."))

        return f"R$ {val:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception:
        return str(valor_raw)


def limpar_texto_pdf(texto):
    if not texto:
        return ""
    return str(texto).encode("latin-1", "ignore").decode("latin-1")



# ============================================================================
# =============================================================================
# 5. GERADORES DE ARQUIVOS (PDF, WORD, EXCEL)
# =============================================================================

# ---------------------------------------------------------------------
# PDF BASE
# ---------------------------------------------------------------------
class PDF(FPDF):
    def header(self):
        self.set_fill_color(50, 50, 50)
        self.rect(0, 0, 210, 30, 'F')
        self.set_font('Arial', 'B', 18)
        self.set_text_color(255, 255, 255)
        self.set_xy(10, 8)
        self.cell(0, 10, 'GEN SYSTEM', 0, 1, 'L')
        self.ln(15)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.set_text_color(128)
        self.cell(0, 10, 'Gerado por Gen System IA.', 0, 0, 'C')


# ---------------------------------------------------------------------
# PDF GENÉRICO (RECIBO / ORÇAMENTO / DECLARAÇÃO)
# ---------------------------------------------------------------------
def criar_pdf(tipo, dados, session_id=None):
    # BLINDAGEM CRÍTICA (evita crash se dados vier inválido)
    dados = dados if isinstance(dados, dict) else {}

    pdf = PDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    dados_limpos = {
        k: limpar_texto_pdf(v)
        for k, v in dados.items()
        if isinstance(v, str)
    }

    if tipo == "recibo":
        pdf.multi_cell(
            0, 8,
            f"RECIBO\n\n"
            f"Valor: {formatar_valor(dados.get('valor'))}\n"
            f"Recebido de: {dados_limpos.get('nome_cliente')}\n"
            f"Referente a: {dados_limpos.get('descricao')}",
            border=1
        )

    elif tipo == "orcamento":
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "ORÇAMENTO", 0, 1, "C")
        pdf.ln(5)
        pdf.set_font("Arial", "", 12)
        pdf.multi_cell(
            0, 8,
            f"Cliente: {dados_limpos.get('cliente')}\n"
            f"Valor Total: {formatar_valor(dados.get('valor'))}"
        )

    elif tipo == "declaracao":
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "DECLARAÇÃO DE CONTEÚDO", 0, 1, "C")
        pdf.ln(5)
        pdf.set_font("Arial", "", 11)

        pdf.multi_cell(
            0, 7,
            f"Remetente: {dados_limpos.get('remetente_nome')}\n"
            f"Documento: {dados_limpos.get('remetente_doc')}\n\n"
            f"Destinatário: {dados_limpos.get('destinatario_nome')}\n"
            f"Documento: {dados_limpos.get('destinatario_doc')}"
        )

    nome = f"{tipo}_{datetime.now().strftime('%H%M%S')}.pdf"
    pdf.output(os.path.join(PASTA_DOCS, nome))

    if session_id:
        salvar_documento_db(session_id, nome, "PDF")

    return nome


# ---------------------------------------------------------------------
# WORD – UTILIDADE VISUAL PADRÃO
# ---------------------------------------------------------------------
def _configurar_documento_word(doc, titulo):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.text = f"Gen System • {datetime.now().strftime('%d/%m/%Y')}"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)

    t = doc.add_paragraph(titulo)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(15)
    doc.add_paragraph("")


# ---------------------------------------------------------------------
# WORD – DECLARAÇÃO
# ---------------------------------------------------------------------
def criar_word_declaracao(dados, session_id):
    doc = Document()
    _configurar_documento_word(doc, "DECLARAÇÃO DE CONTEÚDO")

    doc.add_heading("REMETENTE", level=2)
    doc.add_paragraph(f"{dados.get('remetente_nome')}\n{dados.get('remetente_doc')}")

    doc.add_heading("DESTINATÁRIO", level=2)
    doc.add_paragraph(f"{dados.get('destinatario_nome')}\n{dados.get('destinatario_doc')}")

    doc.add_heading("ITENS", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Descrição"
    table.rows[0].cells[1].text = "Qtd"
    table.rows[0].cells[2].text = "Valor (R$)"

    total = 0.0
    for item in dados.get("lista_itens", []):
        row = table.add_row().cells
        row[0].text = str(item.get("item", ""))
        row[1].text = str(item.get("qtd", 1))
        val = float(item.get("custo", 0))
        row[2].text = f"{val:,.2f}".replace(".", ",")
        total += val * int(item.get("qtd", 1))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"TOTAL DECLARADO: R$ {total:,.2f}".replace(".", ",")).bold = True

    doc.add_paragraph("\n_________________________________\nAssinatura do Remetente")

    nome = f"Declaracao_{datetime.now().strftime('%H%M%S')}.docx"
    doc.save(os.path.join(PASTA_DOCS, nome))

    if session_id:
        salvar_documento_db(session_id, nome, "DECLARACAO_DOCX")

    return nome


# ---------------------------------------------------------------------
# WORD – CONTRATO
# ---------------------------------------------------------------------
def criar_word(tipo, dados, session_id):
    doc = Document()
    _configurar_documento_word(doc, "CONTRATO DE PRESTAÇÃO DE SERVIÇOS")

    doc.add_paragraph(f"CONTRATANTE: {dados.get('contratante')}")
    doc.add_paragraph(f"CONTRATADO: {dados.get('contratado')}")

    doc.add_heading("OBJETO", level=2)
    doc.add_paragraph(dados.get("objeto", ""))

    doc.add_heading("VALOR", level=2)
    doc.add_paragraph(f"R$ {dados.get('valor')}")

    doc.add_paragraph("\nContratante: __________________________")
    doc.add_paragraph("Contratado: __________________________")

    nome = f"Contrato_{datetime.now().strftime('%H%M%S')}.docx"
    doc.save(os.path.join(PASTA_DOCS, nome))

    if session_id:
        salvar_documento_db(session_id, nome, "CONTRATO_DOCX")

    return nome


# ---------------------------------------------------------------------
# WORD – ORDEM DE SERVIÇO
# ---------------------------------------------------------------------
def criar_word_os(dados, session_id):
    doc = Document()
    _configurar_documento_word(doc, "ORDEM DE SERVIÇO")

    doc.add_paragraph(f"CLIENTE: {dados.get('cliente')}")
    doc.add_paragraph(f"EQUIPAMENTO: {dados.get('equipamento')}")
    doc.add_paragraph(f"DEFEITO: {dados.get('defeito')}")

    doc.add_paragraph("\n_________________________________\nAssinatura do Cliente")
    doc.add_paragraph("\n_________________________________\nAssinatura do Técnico")

    nome = f"OS_{datetime.now().strftime('%H%M%S')}.docx"
    doc.save(os.path.join(PASTA_DOCS, nome))

    if session_id:
        salvar_documento_db(session_id, nome, "OS_DOCX")

    return nome


# ---------------------------------------------------------------------
# PDF – ORDEM DE SERVIÇO
# ---------------------------------------------------------------------
def criar_pdf_os(dados, session_id):
    pdf = PDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "ORDEM DE SERVIÇO", 0, 1, "C")
    pdf.ln(5)

    pdf.set_font("Arial", "", 12)
    pdf.multi_cell(
        0, 8,
        f"CLIENTE: {dados.get('cliente')}\n"
        f"EQUIPAMENTO: {dados.get('equipamento')}\n"
        f"DEFEITO: {dados.get('defeito')}"
    )

    nome = f"OS_{datetime.now().strftime('%H%M%S')}.pdf"
    pdf.output(os.path.join(PASTA_DOCS, nome))

    if session_id:
        salvar_documento_db(session_id, nome, "OS_PDF")

    return nome


# ---------------------------------------------------------------------
# EXCEL – UTILIDADES VISUAIS
# ---------------------------------------------------------------------
def _formatar_cabecalho(ws):
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="E0E0E0")
        cell.alignment = Alignment(horizontal="center")
    ws.freeze_panes = "A2"


def _ajustar_colunas(ws):
    for col in ws.columns:
        tamanho = max(len(str(c.value)) if c.value else 0 for c in col)
        ws.column_dimensions[get_column_letter(col[0].column)].width = tamanho + 3


# ---------------------------------------------------------------------
# EXCEL – PRECIFICAÇÃO
# ---------------------------------------------------------------------
def criar_excel_precificacao(dados, session_id):
    # BLINDAGEM CRÍTICA (evita AttributeError / crash)
    dados = dados if isinstance(dados, dict) else {}

    wb = Workbook()
    ws = wb.active
    ws.title = "Precificação"
    ws.append(["Produto", "Custo", "Margem (%)", "Preço Final"])

    for i, item in enumerate(dados.get("itens", [("Exemplo", 10, 100)]), start=2):
        ws.append([
            item[0],
            item[1],
            item[2],
            f"=B{i}+(B{i}*(C{i}/100))"
        ])

    _formatar_cabecalho(ws)
    _ajustar_colunas(ws)

    nome = f"precificacao_{datetime.now().strftime('%H%M%S')}.xlsx"
    wb.save(os.path.join(PASTA_DOCS, nome))

    if session_id:
        salvar_documento_db(session_id, nome, "EXCEL")

    return nome


#----------------------------------------------------------------------
#crias excel
# ---------------------------------------------------------------------
# EXCEL – PLANILHA SIMPLES (FALLBACK / COMPATIBILIDADE)
# ---------------------------------------------------------------------
def criar_excel_simples(dados, tipo=None, session_id=None):
    wb = Workbook()
    ws = wb.active

    # Título seguro
    nome_planilha = str(tipo)[:30] if tipo else "Planilha"
    ws.title = nome_planilha

    # Blindagem total de dados
    if isinstance(dados, dict):
        ws.append(list(dados.keys()))
        ws.append(list(dados.values()))

    elif isinstance(dados, list):
        for item in dados:
            if isinstance(item, (list, tuple)):
                ws.append(item)
            else:
                ws.append([item])

    else:
        ws.append(["Valor"])
        ws.append([str(dados)])

    _formatar_cabecalho(ws)
    _ajustar_colunas(ws)

    # Nome de arquivo seguro
    prefixo = nome_planilha.lower().replace(" ", "_")
    nome = f"{prefixo}_{datetime.now().strftime('%H%M%S')}.xlsx"
    wb.save(os.path.join(PASTA_DOCS, nome))

    if session_id:
        salvar_documento_db(session_id, nome, "EXCEL_SIMPLES")

    return nome


# ---------------------------------------------------------------------
# EXCEL – CAIXA
# ---------------------------------------------------------------------
def criar_excel_caixa(session_id):
    wb = Workbook()
    ws = wb.active
    ws.title = "Fluxo de Caixa"
    ws.append(["Data", "Entrada", "Saída", "Saldo"])
    ws.append(["Hoje", 0, 0, "=B2-C2"])

    for r in range(3, 100):
        ws[f"D{r}"] = f"=D{r-1}+B{r}-C{r}"

    _formatar_cabecalho(ws)
    _ajustar_colunas(ws)

    nome = f"fluxo_caixa_{datetime.now().strftime('%H%M%S')}.xlsx"
    wb.save(os.path.join(PASTA_DOCS, nome))

    if session_id:
        salvar_documento_db(session_id, nome, "EXCEL")

    return nome


# ---------------------------------------------------------------------
# EXCEL – ESTOQUE
# ---------------------------------------------------------------------
def criar_excel_estoque(session_id):
    wb = Workbook()
    ws = wb.active
    ws.title = "Estoque"
    ws.append(["Produto", "Quantidade", "Status"])
    ws.append(["Exemplo", 10, '=IF(B2<=5,"Baixo","OK")'])

    _formatar_cabecalho(ws)
    _ajustar_colunas(ws)

    nome = f"estoque_{datetime.now().strftime('%H%M%S')}.xlsx"
    wb.save(os.path.join(PASTA_DOCS, nome))

    if session_id:
        salvar_documento_db(session_id, nome, "EXCEL")

    return nome


# ---------------------------------------------------------------------
# EXCEL – GRÁFICO
# ---------------------------------------------------------------------
def criar_excel_com_grafico(dados_lista_raw, session_id):
    wb = Workbook()
    ws = wb.active
    ws.title = "Análise Visual"

    if not dados_lista_raw:
        dados = [["Item", "Valor"], ["Exemplo", 10]]
    elif isinstance(dados_lista_raw[0], dict):
        chaves = list(dados_lista_raw[0].keys())
        dados = [chaves] + [[item.get(k) for k in chaves] for item in dados_lista_raw]
    else:
        dados = dados_lista_raw

    for linha in dados:
        ws.append(linha)

    pie = PieChart()
    pie.title = "Gráfico de Análise"
    labels = Reference(ws, min_col=1, min_row=2, max_row=len(dados))
    data = Reference(ws, min_col=2, min_row=1, max_row=len(dados))

    pie.add_data(data, titles_from_data=True)
    pie.set_categories(labels)
    ws.add_chart(pie, "E2")

    _formatar_cabecalho(ws)
    _ajustar_colunas(ws)

    nome = f"grafico_{datetime.now().strftime('%H%M%S')}.xlsx"
    wb.save(os.path.join(PASTA_DOCS, nome))

    if session_id:
        salvar_documento_db(session_id, nome, "EXCEL_GRAFICO")

    return nome


## =============================================================================
# 6. ROTAS DA API (ENDPOINTS)
# =============================================================================

@app.post("/salvar_chave")
def salvar_chave_api(dados: ConfigData):
    global API_KEY_CLIENTE
    try:
        genai.configure(api_key=dados.api_key)
        model = genai.GenerativeModel("models/gemini-flash-latest")
        model.generate_content("Teste")
    except Exception:
        return {"status": "erro", "mensagem": "Chave inválida."}

    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump({"api_key": dados.api_key}, f)

    API_KEY_CLIENTE = dados.api_key
    return {"status": "ok"}


@app.get("/verificar_status")
def verificar_status():
    return {"status": "ativo" if API_KEY_CLIENTE else "pendente"}


# ====================================================================
# ROTA MULTIMODAL: RECEBE ARQUIVOS E IMAGENS DO CHAT
# ====================================================================
@app.post("/chat_com_imagem")
async def chat_com_arquivo_endpoint(
    session_id: str = Form(...),
    texto: str = Form("Por favor, analise este arquivo."),
    arquivo: UploadFile = File(...)
):
    """
    Rota que o frontend chama quando tem um anexo.
    Processa Excel, Word, PDF como texto e PNG/JPG como imagem visual.
    """
    try:
        # 1. Extrai o conteúdo do arquivo usando sua função nativa
        resultado_extracao = await ler_arquivo_para_texto(arquivo)
        
        prompt_final = texto
        imagem_bytes = None
        mime_type = "image/jpeg"

        # 2. Prepara o envio para a IA dependendo do tipo do arquivo
        if resultado_extracao["tipo"] == "imagem":
            imagem_bytes = resultado_extracao["conteudo"]
            mime_type = resultado_extracao["mime"]
            prompt_final = f"Usuário enviou uma imagem. Comando: {texto}"
            
        elif resultado_extracao["tipo"] == "texto":
            prompt_final = f"O usuário enviou um documento com os seguintes dados extraídos:\n\n{resultado_extracao['conteudo']}\n\nComando do usuário: {texto}"
            
        else:
            return {"resposta_gen": f"⚠️ Não consegui ler o formato deste arquivo. Erro: {resultado_extracao.get('conteudo')}"}

        # 3. Salva no histórico do banco de dados
        salvar_mensagem(session_id, "user", f"{texto} [Anexo: {arquivo.filename}]")

        # 4. Aciona a IA com o modo Geral para interpretar o documento/imagem
        prompt_completo = f"""
        {PROMPTS_MODOS['geral']}
        
        AÇÃO SOLICITADA:
        {prompt_final}
        """
        
        raw_response = gerar_com_router(prompt_completo, imagem_bytes, mime_type)
        
        # 5. Limpa a resposta para garantir que o JSON não quebre o chat
        raw_response = re.sub(r"```json|```", "", raw_response).strip()
        try:
            js = json.loads(raw_response)
            resposta_final = js.get("resposta_usuario", raw_response)
        except Exception:
            resposta_final = raw_response

        salvar_mensagem(session_id, "model", resposta_final)
        
        return {"resposta_gen": resposta_final}

    except Exception as e:
        print("[ERRO NO UPLOAD/CHAT MULTIMODAL]:", e)
        return {"resposta_gen": "⚠️ Ocorreu um erro ao processar o seu arquivo."}


# ====================================================================
# ROTA QUE RECEBE OS DADOS DOS FORMULÁRIOS DA BARRA LATERAL
# ====================================================================
@app.post("/gerar_formulario")
def gerar_formulario_endpoint(dados_form: DadosFormulario):
    """
    Rota para receber os dados dos formulários HTML da barra lateral 
    e gerar o documento direto, sem passar pelo chat da IA.
    """
    try:
        tipo = dados_form.tipo.lower()
        formato = dados_form.formato.lower() if dados_form.formato else "pdf"
        dados = dados_form.dados
        session_id = dados_form.session_id
        
        nome_arquivo = ""
        
        # --- Lógica da Ordem de Serviço (OS) ---
        if tipo == "os":
            if formato == "word" or formato == "docx":
                nome_arquivo = criar_word_os(dados, session_id)
            else:
                nome_arquivo = criar_pdf_os(dados, session_id)
        
        # --- Lógica de Recibo / Orçamento ---
        elif tipo in ["recibo", "orcamento"]:
            nome_arquivo = criar_pdf(tipo, dados, session_id)
            
        # --- Lógica de Contrato / Declaração ---
        elif tipo == "contrato":
            nome_arquivo = criar_word("contrato", dados, session_id)
        elif tipo == "declaracao":
            nome_arquivo = criar_word_declaracao(dados, session_id)
            
        else:
            return {"erro": "Tipo de documento não suportado pelo formulário."}
            
        # Retorna o arquivo gerado para o JavaScript fazer o download
        return {"status": "ok", "arquivo": nome_arquivo}
        
    except Exception as e:
        print("[ERRO AO GERAR PELO FORMULARIO]:", e)
        return {"erro": str(e)}
# ====================================================================


@app.get("/baixar_doc/{nome_arquivo}")
def baixar_doc(nome_arquivo: str):
    caminho = os.path.join(PASTA_DOCS, nome_arquivo)
    if not os.path.exists(caminho):
        raise HTTPException(status_code=404, detail="Arquivo não encontrado.")
    return FileResponse(caminho, filename=nome_arquivo)


@app.get("/meus_arquivos/{session_id}")
def listar_arquivos_usuario(session_id: str):
    conn = sqlite3.connect(DB_FILE)
    try:
        rows = conn.execute(
            "SELECT nome_arquivo, tipo, criado_em FROM documentos WHERE session_id=? ORDER BY id DESC LIMIT 20",
            (session_id,)
        ).fetchall()
        return [{"nome": r[0], "tipo": r[1], "data": r[2]} for r in rows]
    finally:
        conn.close()


@app.get("/sessions")
def listar_conversas():
    conn = sqlite3.connect(DB_FILE)
    cur = conn.cursor()
    cur.execute("SELECT session_id, titulo FROM sessoes ORDER BY rowid DESC")
    rows = cur.fetchall()
    conn.close()
    return [{"id": r[0], "titulo": r[1]} for r in rows]


@app.get("/historico/{session_id}")
def carregar_historico(session_id: str):
    conn = sqlite3.connect(DB_FILE)
    rows = conn.execute(
        "SELECT role, content FROM mensagens WHERE session_id=? ORDER BY id ASC",
        (session_id,)
    ).fetchall()
    conn.close()
    return [{"role": r[0], "content": r[1]} for r in rows]


@app.delete("/chat/{session_id}")
def deletar_chat(session_id: str):
    conn = sqlite3.connect(DB_FILE)
    conn.execute("DELETE FROM mensagens WHERE session_id=?", (session_id,))
    conn.execute("DELETE FROM sessoes WHERE session_id=?", (session_id,))
    conn.commit()
    conn.close()
    return {"status": "ok"}

# =============================================================================
# CHAT PRINCIPAL (INTEGRADO AO MOTOR DE DECISÃO)
# =============================================================================

@app.post("/chat")
def conversar_com_gen(pedido: Pedido):
    try:
        # Salva a mensagem do usuário
        salvar_mensagem(pedido.session_id, "user", pedido.texto)

        # 1. OTIMIZAÇÃO: Verifica a intenção ANTES de executar
        # Usamos o motor para decidir se é uma ação de arquivo (rápido, sem custo de API)
        decisao = motor_decisao(pedido.texto)

        if decisao["gerar_arquivo"]:
            # Só chama o executor se realmente for para gerar um arquivo
            resultado = executar_decisao_ia(
                texto_usuario=pedido.texto,
                session_id=pedido.session_id,
                contexto_extra={}
            )

            # Se gerou o arquivo com sucesso, salva e retorna aqui mesmo
            if "arquivo" in resultado:
                salvar_mensagem(pedido.session_id, "model", resultado["resposta_usuario"])
                return resultado

        # 2. CONTEXTO AVANÇADO (MODOS)
        # Se não foi gerado arquivo, segue para a conversa inteligente com contexto
        historico = get_historico_db(pedido.session_id)
        contexto_sql = buscar_dados_tecnicos(pedido.texto)
        dados_mercado, _ = ExternalDataService.get_market_data()

        modo = pedido.modo.lower().strip()
        modo = {
            "jurídico": "juridico",
            "juridico": "juridico",
            "financeiro": "financeiro",
            "finanças": "financeiro",
            "marketing": "marketing",
            "viabilidade": "viabilidade"
        }.get(modo, "geral")

        contexto_web = ""
        if modo in ("juridico", "financeiro", "viabilidade"):
            contexto_web = DeepSearchService.buscar_viabilidade(pedido.texto)

        prompt = f"""
{PROMPTS_MODOS[modo]}
CONTEXTO SQL:
{contexto_sql}

CONTEXTO WEB:
{contexto_web}

DADOS DE MERCADO:
{dados_mercado}

HISTÓRICO:
{historico}

Responda APENAS em JSON.
"""

        raw = gerar_com_router(prompt)
        raw = re.sub(r"```json|```", "", raw).strip()

        try:
            js = json.loads(raw)
        except Exception:
            # Fallback caso a IA não retorne JSON puro
            js = {"resposta_usuario": raw}

        resposta = js.get("resposta_usuario", "Não consegui responder.")
        salvar_mensagem(pedido.session_id, "model", resposta)
        return {"resposta_gen": resposta}

    except Exception as e:
        print("[ERRO CHAT]:", e)
        return {"erro": "Erro interno no chat"}
    
# ============================================================================
# INICIALIZAÇÃO DO APLICATIVO (JANELA E DASHBOARD)
# ============================================================================

def iniciar_dashboard():
    dash = os.path.join(DIRETORIO_EXECUCAO, "dashboard.py")
    if os.path.exists(dash): 
        print("[INFO] Iniciando Dashboard na porta 8501...")
        subprocess.Popen([sys.executable, "-m", "streamlit", "run", dash, "--server.port=8501", "--server.headless=true"], cwd=DIRETORIO_EXECUCAO, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    else:
        print("[AVISO] Dashboard.py nao encontrado na pasta.")

if __name__ == '__main__':
    # 1. Inicia a API (FastAPI) em segundo plano
    t = threading.Thread(target=lambda: uvicorn.run(app, host="127.0.0.1", port=8000, log_level="error"))
    t.daemon = True
    t.start()
    
    # 2. Inicia o painel de métricas (Streamlit)
    iniciar_dashboard()
    time.sleep(1) # Dá 1 segundo para o servidor respirar
    
    # 3. Abre a Janela Principal do Aplicativo
    webview.create_window("Gen System - Dashboard Corporativo", "http://127.0.0.1:8000", width=1200, height=800, resizable=True)
    webview.start()
